import asyncio
import os
import re

from crawl4ai import AsyncWebCrawler, BrowserConfig, CrawlerRunConfig
from crawl4ai.deep_crawling import BFSDeepCrawlStrategy
from crawl4ai.deep_crawling.filters import FilterChain, DomainFilter
from docx import Document

BASE_URL = "https://www.vimalacollege.edu.in"
EXPORT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "export")
MAX_PAGES = 100
MAX_DEPTH = 3


def build_docx(pages: list[dict]) -> Document:
    """Build a single .docx from all scraped pages."""
    doc = Document()
    doc.add_heading("Vimala College — Full Site Export", level=0)

    for i, page in enumerate(pages):
        if i > 0:
            doc.add_page_break()

        doc.add_heading(page["title"], level=1)
        doc.add_paragraph(page["url"], style="Intense Quote")

        for line in page["markdown"].split("\n"):
            line = line.rstrip()
            if not line:
                continue

            heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
            if heading_match:
                level = min(len(heading_match.group(1)) + 1, 4)
                doc.add_heading(heading_match.group(2), level=level)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

    return doc


async def crawl():
    os.makedirs(EXPORT_DIR, exist_ok=True)

    browser_config = BrowserConfig(headless=True, verbose=False)

    strategy = BFSDeepCrawlStrategy(
        max_depth=MAX_DEPTH,
        include_external=False,
        max_pages=MAX_PAGES,
        filter_chain=FilterChain([
            DomainFilter(
                allowed_domains=["vimalacollege.edu.in", "www.vimalacollege.edu.in"],
            ),
        ]),
    )

    run_config = CrawlerRunConfig(
        deep_crawl_strategy=strategy,
        stream=True,
        excluded_tags=["nav", "footer", "header", "script", "style", "noscript", "iframe"],
    )

    pages = []
    seen_urls = set()

    print(f"Deep crawling {BASE_URL} (max {MAX_PAGES} pages, depth {MAX_DEPTH})")
    print(f"Export directory: {EXPORT_DIR}\n")

    async with AsyncWebCrawler(config=browser_config) as crawler:
        async for result in await crawler.arun(BASE_URL, config=run_config):
            if not result.success:
                print(f"  [FAIL] {result.url} — {result.error_message}")
                continue

            normalized = result.url.rstrip("/").lower()
            if normalized in seen_urls:
                continue
            seen_urls.add(normalized)

            markdown = (result.markdown or "").strip()
            if not markdown:
                print(f"  [SKIP] {result.url} — no content")
                continue

            title = result.metadata.get("title", "Untitled") if result.metadata else "Untitled"
            depth = result.metadata.get("depth", "?") if result.metadata else "?"

            pages.append({
                "url": result.url,
                "title": title,
                "markdown": markdown,
            })

            print(f"  [{len(pages):>3}] depth={depth} {result.url}")

    print(f"\nCrawled {len(pages)} pages total.")

    if not pages:
        print("No pages scraped. Exiting.")
        return

    pages.sort(key=lambda p: p["url"])

    # --- Export single combined Markdown ---
    md_path = os.path.join(EXPORT_DIR, "all.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# Vimala College — Full Site Export\n\n")
        for i, page in enumerate(pages):
            if i > 0:
                f.write("\n\n---\n\n")
            f.write(f"## {page['title']}\n\n")
            f.write(f"**Source:** {page['url']}\n\n")
            f.write(page["markdown"])
            f.write("\n")

    # --- Export single combined DOCX ---
    docx_path = os.path.join(EXPORT_DIR, "all.docx")
    doc = build_docx(pages)
    doc.save(docx_path)

    print(f"\nExported:")
    print(f"  {md_path}")
    print(f"  {docx_path}")


def main():
    asyncio.run(crawl())


if __name__ == "__main__":
    main()
